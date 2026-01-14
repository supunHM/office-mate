"""
Flask API for Document Upload with OCR, NLP, and ML Classification
POST /api/documents - Upload and process documents
GET /api/documents - Search and filter documents with pagination
"""
import os
from datetime import datetime, timedelta
from werkzeug.utils import secure_filename
from flask import Blueprint, request, jsonify
from flask_models import db, Document, Tag
from flask_auth import get_current_user
import PyPDF2
from PIL import Image
import pytesseract
import io
import spacy
import joblib
from sklearn.feature_extraction.text import TfidfVectorizer

# Create Blueprint
documents_bp = Blueprint('documents', __name__)

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'tiff', 'bmp', 'doc', 'docx'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
MODEL_PATH = 'models_store/classifier.joblib'

# Ensure upload directory exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Load spaCy model (load once at startup)
try:
    nlp = spacy.load('en_core_web_sm')
except OSError:
    # Download if not found: python -m spacy download en_core_web_sm
    print("Warning: spaCy model 'en_core_web_sm' not found. Please run: python -m spacy download en_core_web_sm")
    nlp = None

# Load ML classifier (if available)
try:
    classifier_data = joblib.load(MODEL_PATH)
    vectorizer = classifier_data['vectorizer']
    classifier = classifier_data['model']
except Exception:
    print("Warning: ML classifier not found. Using fallback categorization.")
    vectorizer = None
    classifier = None


def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_stream):
    """Extract text from PDF using PyPDF2"""
    try:
        pdf_reader = PyPDF2.PdfReader(file_stream)
        text = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return '\n'.join(text)
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""


def extract_text_from_image(file_bytes):
    """Extract text from image using Tesseract OCR"""
    try:
        image = Image.open(io.BytesIO(file_bytes))
        print(f"Image opened: size={image.size}, mode={image.mode}")
        
        # Use Tesseract with English (add '+sin' for Sinhala if installed)
        text = pytesseract.image_to_string(image, lang='eng')
        
        print(f"OCR extraction completed: {len(text)} characters extracted")
        if text:
            print(f"First 100 chars: {text[:100]}")
        
        return text
    except Exception as e:
        print(f"OCR error: {e}")
        import traceback
        traceback.print_exc()
        return ""


def detect_actual_file_type(file_bytes, filename):
    """
    Detect the actual file type by inspecting file content
    Returns: tuple (actual_type, confidence)
    """
    if len(file_bytes) < 4:
        return 'empty', 'high'
    
    # Check file signatures (magic numbers)
    if file_bytes[:4] == b'PK\x03\x04' or file_bytes[:2] == b'PK':
        return 'zip_based', 'high'  # Could be docx, xlsx, etc
    elif file_bytes[:4] == b'%PDF':
        return 'pdf', 'high'
    elif file_bytes[:2] == b'\xff\xd8':
        return 'jpeg', 'high'
    elif file_bytes[:8] == b'\x89PNG\r\n\x1a\n':
        return 'png', 'high'
    elif file_bytes[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
        return 'doc_old', 'high'  # Old .doc format (OLE2)
    
    # Try to decode as UTF-8 text
    try:
        text = file_bytes.decode('utf-8')
        # Check if it's mostly printable
        printable_ratio = sum(c.isprintable() or c in '\n\r\t' for c in text) / len(text)
        if printable_ratio > 0.9:
            return 'text_utf8', 'high'
    except UnicodeDecodeError:
        pass
    
    # Try Latin-1 (common for old files)
    try:
        text = file_bytes.decode('latin-1', errors='ignore')
        printable_ratio = sum(c.isprintable() or c in '\n\r\t' for c in text) / max(len(text), 1)
        if printable_ratio > 0.7:
            return 'text_latin1', 'medium'
    except:
        pass
    
    return 'unknown', 'none'


def extract_text_from_file(file):
    """Extract text from uploaded file based on type"""
    filename = file.filename.lower()
    
    # Read file content and reset pointer
    file.seek(0)  # Ensure we're at the start
    file_bytes = file.read()
    file.seek(0)  # Reset for later reads
    
    print(f"File size: {len(file_bytes)} bytes")
    
    # Check if file is actually empty
    if len(file_bytes) == 0:
        print("Error: File is empty (0 bytes)")
        return ""
    
    # Detect actual file type
    actual_type, confidence = detect_actual_file_type(file_bytes, filename)
    print(f"File type detection: {actual_type} (confidence: {confidence})")
    
    # Handle mismatched file types
    if filename.endswith('.docx') and actual_type == 'text':
        print("Warning: File has .docx extension but is actually plain text")
        try:
            text_content = file_bytes.decode('utf-8', errors='ignore')
            return text_content
        except:
            pass
    
    if filename.endswith('.docx') and actual_type == 'doc_old':
        print("Warning: File has .docx extension but is old .doc format")
        return ""  # Cannot process without antiword or other converter
    
    # Try PDF extraction
    if filename.endswith('.pdf'):
        file_stream = io.BytesIO(file_bytes)
        text = extract_text_from_pdf(file_stream)
        
        # If PDF text extraction failed, try OCR (scanned PDF)
        if not text or len(text.strip()) < 50:
            print("PDF text extraction minimal, trying OCR...")
            text = extract_text_from_image(file_bytes)
        
        return text
    
    # Image files - use OCR
    elif any(filename.endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']):
        return extract_text_from_image(file_bytes)
    
    # Word documents (basic support)
    elif filename.endswith('.docx'):
        # First check if it's actually a valid DOCX (ZIP-based)
        if actual_type != 'zip_based':
            print(f"File has .docx extension but is actually: {actual_type}")
            
            # Try as plain text
            if actual_type in ('text', 'text_lossy'):
                try:
                    text_content = file_bytes.decode('utf-8', errors='ignore')
                    if text_content and len(text_content.strip()) > 50:
                        print("Successfully extracted as plain text")
                        return text_content
                except Exception as e:
                    print(f"Text extraction failed: {e}")
            
            print("Cannot process: File is not a valid .docx format")
            print("Suggestion: Open in Word and save as 'Word Document (.docx)'")
            return ""
        
        # It's a valid ZIP-based file, try to extract as DOCX
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            if text and text.strip():
                print(f"DOCX extraction successful: {len(text)} characters")
                return text
            print("DOCX file has no text content in paragraphs")
            return ""
        except Exception as e:
            error_type = type(e).__name__
            print(f"DOCX extraction error ({error_type}): {e}")
            return ""
    
    # Older .doc format (not supported, need additional library)
    elif filename.endswith('.doc'):
        if actual_type == 'doc_old':
            print("Error: Old .doc format not supported")
            print("Please open in Word and save as .docx format")
        else:
            print(f"File has .doc extension but is actually: {actual_type}")
        return ""
    
    # Fallback: try as plain text for any unknown format
    if actual_type in ('text', 'text_lossy'):
        try:
            text_content = file_bytes.decode('utf-8', errors='ignore')
            if text_content and len(text_content.strip()) > 50:
                print(f"Fallback: Extracted as plain text ({len(text_content)} chars)")
                return text_content
        except:
            pass
    
    return ""


def preprocess_text_with_spacy(text):
    """Preprocess text using spaCy - lemmatization, stopword removal"""
    if not nlp or not text:
        return text
    
    try:
        doc = nlp(text)
        # Lemmatize and remove stopwords, punctuation
        tokens = [token.lemma_.lower() for token in doc 
                 if not token.is_stop and not token.is_punct and token.is_alpha]
        return ' '.join(tokens)
    except Exception as e:
        print(f"spaCy preprocessing error: {e}")
        return text


def classify_document(text):
    """Classify document into category using ML model"""
    if not text or not classifier or not vectorizer:
        return fallback_classification(text)
    
    try:
        # Vectorize text
        text_vectorized = vectorizer.transform([text])
        # Predict category
        category = classifier.predict(text_vectorized)[0]
        return category
    except Exception as e:
        print(f"Classification error: {e}")
        return fallback_classification(text)


def fallback_classification(text):
    """Simple keyword-based classification fallback"""
    if not text:
        return 'unknown'
    
    text_lower = text.lower()
    
    # Finance keywords
    if any(word in text_lower for word in ['invoice', 'payment', 'receipt', 'finance', 'budget', 'cost', 'expense']):
        return 'Finance'
    
    # HR keywords
    if any(word in text_lower for word in ['employee', 'hr', 'leave', 'salary', 'hiring', 'recruitment', 'resignation']):
        return 'HR'
    
    # Procurement keywords
    if any(word in text_lower for word in ['purchase', 'procurement', 'vendor', 'supplier', 'order', 'quotation']):
        return 'Procurement'
    
    # Maintenance keywords
    if any(word in text_lower for word in ['maintenance', 'repair', 'fix', 'broken', 'service', 'equipment']):
        return 'Maintenance'
    
    return 'unknown'


def extract_tags_from_text(text):
    """Extract tags using spaCy named entities and keywords"""
    if not nlp or not text:
        return extract_simple_tags(text)
    
    try:
        doc = nlp(text)
        tags = set()
        
        # Extract named entities
        for ent in doc.ents:
            if ent.label_ in ['ORG', 'PERSON', 'DATE', 'MONEY', 'GPE']:
                tags.add(ent.text.strip())
        
        # Extract key noun phrases
        for chunk in doc.noun_chunks:
            if len(chunk.text.split()) <= 3:  # Max 3 words
                tags.add(chunk.text.strip())
        
        # Limit to top 5 tags
        return list(tags)[:5]
    except Exception as e:
        print(f"Tag extraction error: {e}")
        return extract_simple_tags(text)


def extract_simple_tags(text):
    """Simple keyword-based tag extraction fallback"""
    if not text:
        return []
    
    # Split into words and get most common
    words = text.lower().split()
    # Filter words (length > 3, alphanumeric)
    words = [w.strip('.,!?;:()[]{}') for w in words if len(w) > 3 and w.isalnum()]
    
    # Count frequency
    word_freq = {}
    for word in words:
        word_freq[word] = word_freq.get(word, 0) + 1
    
    # Get top 5 most frequent
    top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:5]
    return [word for word, freq in top_words]


def generate_summary(text, max_sentences=3):
    """Generate a simple summary from text"""
    if not text:
        return ""
    
    # Split into sentences
    sentences = text.replace('\n', ' ').split('.')
    sentences = [s.strip() for s in sentences if s.strip()]
    
    # Return first N sentences
    summary_sentences = sentences[:max_sentences]
    return '. '.join(summary_sentences) + '.' if summary_sentences else text[:200]


def get_or_create_tag(tag_name):
    """Get existing tag or create new one"""
    tag = Tag.query.filter_by(name=tag_name).first()
    if not tag:
        tag = Tag(name=tag_name)
        db.session.add(tag)
    return tag


@documents_bp.route('/api/documents', methods=['POST'])
def upload_document():
    """
    Upload and process document
    
    Request: multipart/form-data with 'file' field
    Response: JSON with document details
    """
    # Check if file is in request
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    # Check if filename is empty
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    # Validate file type
    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not allowed'}), 400
    
    # Get authenticated user
    user_id = get_current_user()
    if not user_id:
        return jsonify({'error': 'Authentication required'}), 401
    
    try:
        # Secure filename
        original_filename = secure_filename(file.filename)
        
        # Create user-specific upload directory
        user_upload_dir = os.path.join(UPLOAD_FOLDER, str(user_id))
        os.makedirs(user_upload_dir, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{original_filename}"
        file_path = os.path.join(user_upload_dir, filename)
        
        # Save file
        file.save(file_path)
        
        # Step 1: Extract text
        print(f"Extracting text from {original_filename}...")
        extracted_text = extract_text_from_file(file)
        
        # Detect file type for metadata
        file.seek(0)
        file_bytes = file.read()
        file_type, _ = detect_actual_file_type(file_bytes, original_filename)
        
        # For images, be more lenient with minimum text requirement
        min_chars = 5 if original_filename.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')) else 10
        
        if not extracted_text or len(extracted_text.strip()) < min_chars:
            error_msg = 'Could not extract text from document. '
            
            if original_filename.lower().endswith('.doc'):
                error_msg += 'Old .doc format is not supported. Please open in Microsoft Word and Save As → Word Document (.docx).'
            elif original_filename.lower().endswith('.docx'):
                if file_type == 'text':
                    error_msg += 'File appears to be plain text with .docx extension. Please save as actual .docx in Word.'
                elif file_type == 'doc_old':
                    error_msg += 'File is in old .doc format with .docx extension. Open in Word and resave as .docx.'
                else:
                    error_msg += 'File may be corrupted. Try: 1) Open in Word 2) Select All & Copy 3) New document 4) Paste 5) Save as .docx'
            elif original_filename.lower().endswith('.pdf'):
                error_msg += 'PDF has no extractable text. Try a PDF with selectable text or better quality for OCR.'
            else:
                error_msg += 'File format not recognized or has no readable text content.'
            
            return jsonify({'error': error_msg}), 400
        
        # Step 2: Preprocess text with spaCy
        print("Preprocessing text with spaCy...")
        preprocessed_text = preprocess_text_with_spacy(extracted_text)
        
        # Step 3: Classify document
        print("Classifying document...")
        category = classify_document(preprocessed_text)
        
        # Step 4: Extract tags
        print("Extracting tags...")
        tag_names = extract_tags_from_text(extracted_text)
        
        # Step 5: Generate summary
        summary = generate_summary(extracted_text)
        
        # Step 6: Save to database
        print("Saving to database...")
        document = Document(
            file_path=file_path,
            original_name=original_filename,
            text=extracted_text,
            category=category,
            user_id=user_id
        )
        
        # Add tags
        for tag_name in tag_names:
            if tag_name:  # Skip empty tags
                tag = get_or_create_tag(tag_name)
                document.tags.append(tag)
        
        db.session.add(document)
        db.session.commit()
        
        # Step 7: Return response
        response_data = {
            'id': document.id,
            'category': category,
            'tags': tag_names,
            'summary': summary,
            'filename': original_filename,
            'created_at': document.created_at.isoformat(),
            'extracted_text_length': len(extracted_text),
            'extraction_method': file_type  # Shows if it was OCR (image) or direct extraction (pdf/docx)
        }
        
        # Add first 500 characters of extracted text for preview
        if extracted_text:
            response_data['text_preview'] = extracted_text[:500] + ('...' if len(extracted_text) > 500 else '')
        
        return jsonify(response_data), 201
        
    except Exception as e:
        db.session.rollback()
        print(f"Error processing document: {str(e)}")
        return jsonify({'error': f'Failed to process document: {str(e)}'}), 500


# GET endpoint with search and filtering
@documents_bp.route('/api/documents', methods=['GET'])
def get_documents():
    """
    Search and filter documents with pagination
    Query Parameters:
    - q: keyword search in document text or tags
    - category: filter by category (Finance, HR, Procurement, Maintenance)
    - start_date: filter by created_at >= start_date (YYYY-MM-DD)
    - end_date: filter by created_at <= end_date (YYYY-MM-DD)
    - page: page number (default: 1)
    - per_page: items per page (default: 20, max: 100)
    - user_id: user ID (replace with actual auth)
    """
    try:
        # Get query parameters
        search_query = request.args.get('q', '').strip()
        category = request.args.get('category', '').strip()
        start_date = request.args.get('start_date', '').strip()
        end_date = request.args.get('end_date', '').strip()
        page = int(request.args.get('page', 1))
        per_page = min(int(request.args.get('per_page', 20)), 100)
        
        # Get authenticated user
        user_id = get_current_user()
        if not user_id:
            return jsonify({'error': 'Authentication required'}), 401
        
        # Start building query
        query = Document.query.filter_by(user_id=user_id)
        
        # Keyword search in text or tags
        if search_query:
            # Search in document text (case-insensitive)
            search_pattern = f"%{search_query}%"
            text_filter = Document.text.ilike(search_pattern)
            
            # Search in tags (join with Tag table)
            tag_filter = Document.tags.any(Tag.name.ilike(search_pattern))
            
            # Combine with OR
            query = query.filter(db.or_(text_filter, tag_filter))
        
        # Category filter
        if category:
            query = query.filter(Document.category == category)
        
        # Date range filter
        if start_date:
            try:
                start_dt = datetime.strptime(start_date, '%Y-%m-%d')
                query = query.filter(Document.created_at >= start_dt)
            except ValueError:
                return jsonify({'error': 'Invalid start_date format. Use YYYY-MM-DD'}), 400
        
        if end_date:
            try:
                end_dt = datetime.strptime(end_date, '%Y-%m-%d')
                # Add one day to include the entire end_date
                from datetime import timedelta
                end_dt = end_dt + timedelta(days=1)
                query = query.filter(Document.created_at < end_dt)
            except ValueError:
                return jsonify({'error': 'Invalid end_date format. Use YYYY-MM-DD'}), 400
        
        # Ordering: relevance (if search query) then recency
        if search_query:
            # SQLite doesn't have great relevance scoring, but we can prioritize:
            # 1. Exact matches in tags (higher relevance)
            # 2. Then order by recency
            # Use a simple approach: documents with matching tags come first
            query = query.outerjoin(Document.tags).order_by(
                db.case(
                    (Tag.name.ilike(search_pattern), 0),
                    else_=1
                ),
                Document.created_at.desc()
            )
        else:
            # Default: order by recency only
            query = query.order_by(Document.created_at.desc())
        
        # Paginate results
        paginated = query.paginate(page=page, per_page=per_page, error_out=False)
        
        # Format response
        documents = [{
            'id': doc.id,
            'original_name': doc.original_name,
            'category': doc.category,
            'tags': [tag.name for tag in doc.tags],
            'created_at': doc.created_at.isoformat(),
            'text_preview': doc.text[:200] + '...' if len(doc.text) > 200 else doc.text
        } for doc in paginated.items]
        
        return jsonify({
            'documents': documents,
            'pagination': {
                'page': paginated.page,
                'per_page': paginated.per_page,
                'total': paginated.total,
                'pages': paginated.pages,
                'has_next': paginated.has_next,
                'has_prev': paginated.has_prev
            }
        }), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# Optional: GET endpoint for single document
@documents_bp.route('/api/documents/<int:document_id>', methods=['GET'])
def get_document(document_id):
    """Get a specific document"""
    # Get authenticated user
    user_id = get_current_user()
    if not user_id:
        return jsonify({'error': 'Authentication required'}), 401
    
    # Get document and verify ownership
    document = Document.query.filter_by(id=document_id, user_id=user_id).first()
    if not document:
        return jsonify({'error': 'Document not found'}), 404
    
    return jsonify({
        'id': document.id,
        'original_name': document.original_name,
        'category': document.category,
        'text': document.text,
        'tags': [tag.name for tag in document.tags],
        'summary': generate_summary(document.text),
        'created_at': document.created_at.isoformat()
    }), 200


# Debug endpoint to test file extraction
@documents_bp.route('/api/documents/test-extract', methods=['POST'])
def test_extract():
    """
    Test endpoint to extract text from file without saving
    Returns detailed extraction info for debugging
    """
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    try:
        import magic
        has_magic = True
    except:
        has_magic = False
    
    # Get file info
    filename = file.filename
    file_bytes = file.read()
    file.seek(0)
    
    result = {
        'filename': filename,
        'size_bytes': len(file_bytes),
        'size_kb': round(len(file_bytes) / 1024, 2),
        'allowed': allowed_file(filename)
    }
    
    # Detect actual file type
    actual_type, confidence = detect_actual_file_type(file_bytes, filename)
    result['detected_type'] = actual_type
    result['detection_confidence'] = confidence
    
    # Show first bytes (hex) for debugging
    result['file_signature'] = ' '.join(f'{b:02x}' for b in file_bytes[:16])
    
    # Try to detect file type
    if has_magic:
        try:
            mime = magic.Magic(mime=True)
            result['mime_type'] = mime.from_buffer(file_bytes)
        except:
            result['mime_type'] = 'unknown'
    
    # Try extraction
    try:
        extracted_text = extract_text_from_file(file)
        result['extraction_success'] = bool(extracted_text and len(extracted_text.strip()) >= 10)
        result['text_length'] = len(extracted_text) if extracted_text else 0
        result['text_preview'] = extracted_text[:200] if extracted_text else None
        
        if not extracted_text or len(extracted_text.strip()) < 10:
            result['error'] = 'No text extracted or text too short'
            result['suggestion'] = 'Check file format and ensure it contains readable text'
    except Exception as e:
        result['extraction_success'] = False
        result['error'] = str(e)
        result['error_type'] = type(e).__name__
    
    return jsonify(result), 200
